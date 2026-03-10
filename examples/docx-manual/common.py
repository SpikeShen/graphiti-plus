"""
Shared utilities for docx-manual test scripts.

Provides: Graphiti client construction, Word document parsing, clear/ingest helpers.
"""

import logging
import os
import re
import time as _time
from datetime import datetime, timezone
from pathlib import Path

from dotenv import load_dotenv

from graphiti_core import Graphiti
from graphiti_core.cross_encoder.bedrock_reranker_client import BedrockRerankerClient
from graphiti_core.embedder.bedrock_nova import BedrockNovaEmbedder, BedrockNovaEmbedderConfig
from graphiti_core.llm_client.bedrock_client import BedrockLLMClient
from graphiti_core.llm_client.config import LLMConfig
from graphiti_core.logging import create_s3_logger
from graphiti_core.nodes import (
    ContentBlock,
    ContentBlockType,
    SemanticRole,
)
from graphiti_core.vector_store.s3_vectors_client import S3VectorsClient, S3VectorsConfig

load_dotenv()

os.environ.setdefault('GRAPHITI_LLM_TRACE', 'false')

logging.getLogger('neo4j').setLevel(logging.WARNING)
logging.getLogger('botocore.credentials').setLevel(logging.WARNING)
logging.getLogger('httpx').setLevel(logging.WARNING)

REGION = os.environ['AWS_REGION']
LLM_MODEL = os.environ['BEDROCK_MODEL']
EMBEDDING_MODEL = os.environ['BEDROCK_EMBEDDING_MODEL']
EMBEDDING_DIM = int(os.environ.get('BEDROCK_EMBEDDING_DIM', '1024'))
LLM_TEMPERATURE = float(os.environ.get('LLM_TEMPERATURE', '0.1'))
S3_VECTORS_BUCKET = os.environ['S3_VECTORS_BUCKET']

DOCX_PATH = os.path.join(os.path.dirname(__file__), 'test-document.docx')
GROUP_ID = 'docx-manual-test'


def build_graphiti() -> Graphiti:
    """Build Graphiti instance with all clients from .env config."""
    llm_client = BedrockLLMClient(
        config=LLMConfig(model=LLM_MODEL, small_model=LLM_MODEL, temperature=LLM_TEMPERATURE),
        region_name=REGION,
    )
    embedder = BedrockNovaEmbedder(
        config=BedrockNovaEmbedderConfig(
            model_id=EMBEDDING_MODEL, region_name=REGION, embedding_dim=EMBEDDING_DIM,
        )
    )
    cross_encoder = BedrockRerankerClient(client=llm_client.client, model_id=LLM_MODEL)
    s3_vectors = S3VectorsClient(
        config=S3VectorsConfig(
            vector_bucket_name=S3_VECTORS_BUCKET, region_name=REGION, embedding_dim=EMBEDDING_DIM,
        )
    )
    return Graphiti(
        os.environ['NEO4J_URI'], os.environ['NEO4J_USER'], os.environ['NEO4J_PASSWORD'],
        llm_client=llm_client, embedder=embedder,
        cross_encoder=cross_encoder, s3_vectors=s3_vectors,
        s3_logger=create_s3_logger(),
    )


# ---------------------------------------------------------------------------
# Word document parser → list[ContentBlock]
# ---------------------------------------------------------------------------

# Namespace constants for docx XML
_WP_NS = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
_A_NS = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
_R_NS = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
_W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def _paragraph_has_image(paragraph) -> bool:
    """Check if a docx paragraph contains an embedded image."""
    el = paragraph._element
    return bool(el.findall(f'.//{_WP_NS}inline') or el.findall(f'.//{_WP_NS}anchor'))


def _extract_image_bytes(paragraph, doc) -> list[tuple[bytes, str]]:
    """Extract (raw_bytes, mime_type) for each image in a paragraph."""
    images = []
    el = paragraph._element
    # Find all blip elements (image references)
    blips = el.findall(
        './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
    )
    for blip in blips:
        embed_id = blip.get(f'{_R_NS}embed')
        if not embed_id:
            continue
        rel = doc.part.rels.get(embed_id)
        if rel and hasattr(rel, 'target_part'):
            part = rel.target_part
            raw = part.blob
            ct = part.content_type or 'image/png'
            images.append((raw, ct))
    return images


def _style_to_semantic_role(style_name: str) -> SemanticRole:
    """Map docx paragraph style to SemanticRole."""
    s = style_name.lower()
    if 'heading 1' in s:
        return SemanticRole.title
    if 'heading' in s:
        return SemanticRole.section_header
    if 'list' in s:
        return SemanticRole.list_item
    if 'caption' in s:
        return SemanticRole.caption
    if 'toc' in s:
        return SemanticRole.header  # table of contents → treat as header/meta
    return SemanticRole.body


def parse_docx(file_path: str) -> list[ContentBlock]:
    """Parse a .docx file into an ordered list of ContentBlocks.

    - Text paragraphs → ContentBlock(block_type=text)
    - Embedded images → ContentBlock(block_type=image, _raw_bytes=...)
    - Consecutive text paragraphs with the same semantic role are merged.
    - TOC entries are skipped.
    - Table content is extracted as text blocks.
    """
    from docx import Document

    doc = Document(file_path)
    blocks: list[ContentBlock] = []
    idx = 0
    pending_text: list[str] = []
    pending_role: SemanticRole = SemanticRole.body

    def _flush_text():
        nonlocal idx
        if not pending_text:
            return
        merged = '\n'.join(pending_text)
        blocks.append(ContentBlock(
            index=idx,
            block_type=ContentBlockType.text,
            text=merged,
            semantic_role=pending_role,
        ))
        idx += 1
        pending_text.clear()

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else 'Normal'
        role = _style_to_semantic_role(style_name)

        # Skip TOC entries
        if 'toc' in style_name.lower():
            continue

        text = para.text.strip()
        has_img = _paragraph_has_image(para)

        # If paragraph has image(s), flush pending text first
        if has_img:
            _flush_text()
            # Also add any text in this paragraph before the image
            if text:
                blocks.append(ContentBlock(
                    index=idx,
                    block_type=ContentBlockType.text,
                    text=text,
                    semantic_role=role,
                ))
                idx += 1

            # Extract images
            for raw_bytes, mime_type in _extract_image_bytes(para, doc):
                fmt = mime_type.split('/')[-1] if '/' in mime_type else 'png'
                block = ContentBlock(
                    index=idx,
                    block_type=ContentBlockType.image,
                    mime_type=mime_type,
                    description=f'操作截图（文档第{idx}个内容块）',
                    metadata={'format': fmt, 'size_bytes': len(raw_bytes)},
                )
                block._raw_bytes = raw_bytes
                blocks.append(block)
                idx += 1
            continue

        # Pure text paragraph
        if not text:
            continue

        # Headings always start a new block
        if role in (SemanticRole.title, SemanticRole.section_header):
            _flush_text()
            blocks.append(ContentBlock(
                index=idx,
                block_type=ContentBlockType.text,
                text=text,
                semantic_role=role,
            ))
            idx += 1
            continue

        # Merge consecutive body/list_item text
        if pending_text and pending_role != role:
            _flush_text()
        pending_role = role
        pending_text.append(text)

    _flush_text()

    # Extract tables
    for ti, table in enumerate(doc.tables):
        rows_text = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows_text.append(' | '.join(cells))
        table_text = '\n'.join(rows_text)
        blocks.append(ContentBlock(
            index=idx,
            block_type=ContentBlockType.table,
            text=table_text,
            semantic_role=SemanticRole.body,
            description=f'表格{ti + 1}',
            metadata={'rows': len(table.rows), 'cols': len(table.columns)},
        ))
        idx += 1

    return blocks


def split_blocks_by_section(blocks: list[ContentBlock]) -> list[tuple[str, list[ContentBlock]]]:
    """Split content blocks into sections by Heading 1/2.

    Returns list of (section_name, blocks) tuples.
    Blocks before the first heading go into '前言'.
    """
    sections: list[tuple[str, list[ContentBlock]]] = []
    current_name = '前言'
    current_blocks: list[ContentBlock] = []

    for block in blocks:
        if (
            block.block_type == ContentBlockType.text
            and block.semantic_role in (SemanticRole.title, SemanticRole.section_header)
        ):
            if current_blocks:
                sections.append((current_name, current_blocks))
            current_name = block.text or '未命名章节'
            current_blocks = [block]
        else:
            current_blocks.append(block)

    if current_blocks:
        sections.append((current_name, current_blocks))

    return sections


async def clear_all(graphiti: Graphiti):
    """Clear Neo4j data and rebuild S3 Vectors indices."""
    print('Clearing all data...', flush=True)
    await graphiti.driver.execute_query("MATCH (n) DETACH DELETE n")
    print('  Neo4j cleared.', flush=True)

    if graphiti.s3_vectors is not None:
        graphiti.s3_vectors.delete_all_indices()
        print('  S3 Vectors indices deleted.', flush=True)

    await graphiti.build_indices_and_constraints()
    print('  Indices rebuilt.', flush=True)


async def ingest_sections(graphiti: Graphiti, sections: list[tuple[str, list[ContentBlock]]]):
    """Ingest document sections via add_document_episode (full multimodal flow).

    Uses add_document_episode() which orchestrates:
    parse → image embedding → S3 upload → vision LLM description →
    build content → image_embedding_map → delegate to add_episode.

    Saga chaining: all sections belong to the same saga (document name),
    linked via NEXT_EPISODE edges for sequential reading order.
    """
    saga_name = os.path.splitext(os.path.basename(DOCX_PATH))[0]
    saga_previous_episode_uuid: str | None = None

    print(f'\n--- Ingesting {len(sections)} sections (saga={saga_name!r}) ---', flush=True)
    total_start = _time.time()
    for i, (name, section_blocks) in enumerate(sections):
        # Re-index blocks within section (0-based)
        for j, block in enumerate(section_blocks):
            block.index = j

        n_text = sum(1 for b in section_blocks if b.block_type == ContentBlockType.text)
        n_img = sum(1 for b in section_blocks if b.block_type == ContentBlockType.image)

        ep_start = _time.time()
        result = await graphiti.add_document_episode(
            name=f'MAP-CCS-Tagging-{i:02d}-{name[:30]}',
            content_blocks=section_blocks,
            source_description='Global MAP 2.0 CCS Tagging 操作指南',
            reference_time=datetime.now(timezone.utc),
            group_id=GROUP_ID,
            saga=saga_name,
            saga_previous_episode_uuid=saga_previous_episode_uuid,
        )
        # Chain: next section links to this episode
        saga_previous_episode_uuid = result.episode.uuid
        print(
            f'  [{i}] {name[:40]:40s} ({n_text} text, {n_img} img) {_time.time() - ep_start:.1f}s',
            flush=True,
        )
    print(f'Total: {_time.time() - total_start:.1f}s', flush=True)


async def print_graph_stats(graphiti: Graphiti):
    """Print entity/edge/episode counts."""
    import json

    print('\n--- Graph Stats ---', flush=True)
    for label, display in [('Entity', 'Entities'), ('Episodic', 'Episodes'), ('Saga', 'Sagas')]:
        records, _, _ = await graphiti.driver.execute_query(
            f"MATCH (n:{label}) WHERE n.group_id = $gid RETURN count(n) AS cnt",
            params={'gid': GROUP_ID},
        )
        print(f'  {display}: {records[0]["cnt"]}', flush=True)

    records, _, _ = await graphiti.driver.execute_query(
        "MATCH ()-[e:RELATES_TO]->() WHERE e.group_id = $gid RETURN count(e) AS cnt",
        params={'gid': GROUP_ID},
    )
    print(f'  Edges: {records[0]["cnt"]}', flush=True)

    records, _, _ = await graphiti.driver.execute_query(
        "MATCH ()-[e:RELATES_TO]->() WHERE e.group_id = $gid "
        "RETURN count(e) AS total, "
        "sum(CASE WHEN e.source_excerpt IS NOT NULL AND e.source_excerpt <> '' THEN 1 ELSE 0 END) AS with_excerpt",
        params={'gid': GROUP_ID},
    )
    r = records[0]
    print(f'  Edges with source_excerpt: {r["with_excerpt"]}/{r["total"]}', flush=True)

    # Saga edge stats (HAS_EPISODE, NEXT_EPISODE)
    for rel_type, display in [('HAS_EPISODE', 'HAS_EPISODE'), ('NEXT_EPISODE', 'NEXT_EPISODE')]:
        records, _, _ = await graphiti.driver.execute_query(
            f"MATCH ()-[r:{rel_type}]->() WHERE r.group_id = $gid RETURN count(r) AS cnt",
            params={'gid': GROUP_ID},
        )
        print(f'  {display}: {records[0]["cnt"]}', flush=True)

    # DescribesEdge stats
    records, _, _ = await graphiti.driver.execute_query(
        "MATCH ()-[d:DESCRIBES]->() WHERE d.group_id = $gid RETURN count(d) AS cnt",
        params={'gid': GROUP_ID},
    )
    print(f'  DescribesEdges: {records[0]["cnt"]}', flush=True)

    # Narrative excerpts stats
    records, _, _ = await graphiti.driver.execute_query(
        "MATCH (e:Episodic) WHERE e.group_id = $gid "
        "RETURN e.name AS name, e.narrative_excerpts AS narrative_excerpts",
        params={'gid': GROUP_ID},
    )
    total_narratives = 0
    for r in records:
        raw = r.get('narrative_excerpts', '[]')
        try:
            narr_list = json.loads(raw) if isinstance(raw, str) else (raw or [])
        except Exception:
            narr_list = []
        total_narratives += len(narr_list)
    print(f'  Narrative excerpts: {total_narratives}', flush=True)

    # content_blocks stats
    records, _, _ = await graphiti.driver.execute_query(
        "MATCH (e:Episodic) WHERE e.group_id = $gid "
        "RETURN e.name AS name, e.content_blocks AS content_blocks",
        params={'gid': GROUP_ID},
    )
    for r in records:
        raw = r.get('content_blocks', '[]')
        try:
            cb_list = json.loads(raw) if isinstance(raw, str) else (raw or [])
        except Exception:
            cb_list = []
        n_text = sum(1 for b in cb_list if b.get('block_type') == 'text')
        n_img = sum(1 for b in cb_list if b.get('block_type') == 'image')
        n_tbl = sum(1 for b in cb_list if b.get('block_type') == 'table')
        print(f'  Episode {r["name"]}: {len(cb_list)} blocks ({n_text} text, {n_img} img, {n_tbl} table)', flush=True)
