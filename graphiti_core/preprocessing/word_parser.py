"""
Word document (.docx) parser.

Extracts text paragraphs and embedded images in reading order,
producing an ordered list of ContentBlock objects.
"""

import logging

from ..nodes import ContentBlock, ContentBlockType, SemanticRole
from .parser import DocumentParser, default_registry

logger = logging.getLogger(__name__)

# XML namespace constants for docx
_WP_NS = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
_A_NS = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
_R_NS = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'


def _paragraph_has_image(paragraph) -> bool:
    """Check if a docx paragraph contains an embedded image."""
    el = paragraph._element
    return bool(el.findall(f'.//{_WP_NS}inline') or el.findall(f'.//{_WP_NS}anchor'))


def _extract_image_bytes(paragraph, doc) -> list[tuple[bytes, str]]:
    """Extract (raw_bytes, content_type) for each image in a paragraph."""
    images = []
    blips = paragraph._element.findall(f'.//{_A_NS}blip')
    for blip in blips:
        embed_id = blip.get(f'{_R_NS}embed')
        if not embed_id:
            continue
        rel = doc.part.rels.get(embed_id)
        if rel and hasattr(rel, 'target_part'):
            part = rel.target_part
            images.append((part.blob, part.content_type or 'image/png'))
    return images


def _style_to_semantic_role(style_name: str) -> SemanticRole:
    """Map docx paragraph style to SemanticRole."""
    s = style_name.lower()
    if 'heading 1' in s:
        return SemanticRole.title
    if 'heading' in s:
        return SemanticRole.section_header
    if 'list' in s:
        return SemanticRole.list_item
    if 'caption' in s:
        return SemanticRole.caption
    if 'toc' in s:
        return SemanticRole.header
    return SemanticRole.body


class WordDocumentParser(DocumentParser):
    """Parse .docx files, extracting text paragraphs and embedded images in order.

    - Consecutive text paragraphs with the same semantic role are merged.
    - TOC entries are skipped.
    - Tables are extracted as text blocks with Markdown-like representation.
    - Image binary data is stored in ContentBlock._raw_bytes (not serialized).
    """

    def supported_extensions(self) -> list[str]:
        return ['.docx']

    async def parse(self, file_path: str) -> list[ContentBlock]:
        from docx import Document

        doc = Document(file_path)
        blocks: list[ContentBlock] = []
        idx = 0
        pending_text: list[str] = []
        pending_role: SemanticRole = SemanticRole.body

        def _flush_text():
            nonlocal idx
            if not pending_text:
                return
            blocks.append(ContentBlock(
                index=idx,
                block_type=ContentBlockType.text,
                text='\n'.join(pending_text),
                semantic_role=pending_role,
            ))
            idx += 1
            pending_text.clear()

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else 'Normal'
            role = _style_to_semantic_role(style_name)

            # Skip TOC entries
            if 'toc' in style_name.lower():
                continue

            text = para.text.strip()
            has_img = _paragraph_has_image(para)

            if has_img:
                _flush_text()
                if text:
                    blocks.append(ContentBlock(
                        index=idx,
                        block_type=ContentBlockType.text,
                        text=text,
                        semantic_role=role,
                    ))
                    idx += 1

                for raw_bytes, mime_type in _extract_image_bytes(para, doc):
                    fmt = mime_type.split('/')[-1] if '/' in mime_type else 'png'
                    block = ContentBlock(
                        index=idx,
                        block_type=ContentBlockType.image,
                        mime_type=mime_type,
                        metadata={'format': fmt, 'size_bytes': len(raw_bytes)},
                    )
                    block._raw_bytes = raw_bytes
                    blocks.append(block)
                    idx += 1
                continue

            if not text:
                continue

            # Headings always start a new block
            if role in (SemanticRole.title, SemanticRole.section_header):
                _flush_text()
                blocks.append(ContentBlock(
                    index=idx,
                    block_type=ContentBlockType.text,
                    text=text,
                    semantic_role=role,
                ))
                idx += 1
                continue

            # Merge consecutive body/list_item text
            if pending_text and pending_role != role:
                _flush_text()
            pending_role = role
            pending_text.append(text)

        _flush_text()

        # Extract tables
        for ti, table in enumerate(doc.tables):
            rows_text = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                rows_text.append(' | '.join(cells))
            blocks.append(ContentBlock(
                index=idx,
                block_type=ContentBlockType.table,
                text='\n'.join(rows_text),
                semantic_role=SemanticRole.body,
                description=f'表格{ti + 1}',
                metadata={'rows': len(table.rows), 'cols': len(table.columns)},
            ))
            idx += 1

        logger.info(
            'Parsed %s: %d blocks (%d text, %d image, %d table)',
            file_path,
            len(blocks),
            sum(1 for b in blocks if b.block_type == ContentBlockType.text),
            sum(1 for b in blocks if b.block_type == ContentBlockType.image),
            sum(1 for b in blocks if b.block_type == ContentBlockType.table),
        )
        return blocks


# Register with default registry
default_registry.register(WordDocumentParser())
